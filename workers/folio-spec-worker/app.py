#!/usr/bin/env python3
"""
Folio-Spec Worker - Converts folio content to NLdoc spec JSON

This worker consumes jobs from station-folio-spec and generates a spec document
that can be converted to HTML/TipTap by downstream workers.

Supports: PDF and DOCX files

Flow:
  folio-spec station → worker.folio-spec.jobs → THIS WORKER → specs.<docId> → station-spec-html → html-writer
"""

import os
import json
import uuid
import pika
import time
from datetime import datetime
from minio import Minio
from io import BytesIO

# PDF text extraction
import fitz  # PyMuPDF

# DOCX text extraction
from docx import Document as DocxDocument
from docx.shared import Pt

# Configuration
# Build AMQP URL from separate env vars if AMQP_URL not provided
AMQP_URL = os.environ.get("AMQP_URL")
if not AMQP_URL:
    AMQP_HOST = os.environ.get("AMQP_HOST", "localhost")
    AMQP_PORT = os.environ.get("AMQP_PORT", "5672")
    AMQP_USERNAME = os.environ.get("AMQP_USERNAME", "guest")
    AMQP_PASSWORD = os.environ.get("AMQP_PASSWORD", "guest")
    AMQP_URL = f"amqp://{AMQP_USERNAME}:{AMQP_PASSWORD}@{AMQP_HOST}:{AMQP_PORT}/"

MINIO_ENDPOINT = os.environ.get("MINIO_ENDPOINT", "minio:9000")
MINIO_USER = os.environ.get("MINIO_USER", os.environ.get("MINIO_ACCESS_KEY", "minio"))
MINIO_PASS = os.environ.get("MINIO_PASS", os.environ.get("MINIO_SECRET_KEY", "minio123"))

WORKER_NAME = "worker-folio-spec"
# IMPORTANT: workerInstance must be stable across container restarts, otherwise station-folio-spec can
# treat worker results as coming from an "unknown worker". In Kubernetes, HOSTNAME == pod name.
INSTANCE_NAME = (
    os.environ.get("KIMI_WORKER_INSTANCE")
    or os.environ.get("HOSTNAME")
    or f"folio-spec-worker-{uuid.uuid4().hex[:8]}"
)

def get_minio():
    return Minio(MINIO_ENDPOINT, access_key=MINIO_USER, secret_key=MINIO_PASS, secure=False)

def extract_text_from_pdf(minio, bucket_name, filename):
    """Extract structured text from each page of a PDF file"""
    print(f"[worker] Extracting text from {bucket_name}/{filename}")
    
    try:
        # Download PDF from MinIO
        response = minio.get_object(bucket_name, filename)
        pdf_bytes = response.read()
        response.close()
        response.release_conn()
        
        # Open PDF with PyMuPDF
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        pages = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extract text blocks with position info for structure detection
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
            
            page_content = []
            for block in blocks:
                if block["type"] == 0:  # Text block
                    for line in block.get("lines", []):
                        line_text = ""
                        font_size = 0
                        is_bold = False
                        
                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            font_size = max(font_size, span.get("size", 12))
                            font_name = span.get("font", "").lower()
                            is_bold = is_bold or "bold" in font_name
                            line_text += text
                        
                        line_text = line_text.strip()
                        if line_text:
                            # Fix encoding issues (Windows-1252 artifacts in UTF-8)
                            line_text = fix_encoding(line_text)
                            
                            # Detect structure based on font size and style
                            if font_size >= 16 or is_bold:
                                page_content.append({"type": "heading", "level": 1 if font_size >= 18 else 2, "text": line_text})
                            else:
                                page_content.append({"type": "paragraph", "text": line_text})
            
            pages.append({
                "page_number": page_num + 1,
                "content": page_content
            })
            print(f"[worker] Page {page_num + 1}: {len(page_content)} blocks")
        
        doc.close()
        return pages
        
    except Exception as e:
        print(f"[worker] Error extracting PDF text: {e}")
        import traceback
        traceback.print_exc()
        return None

def extract_runs_with_formatting(paragraph):
    """Extract text from paragraph runs, preserving inline formatting"""
    if not paragraph.runs:
        return [{"type": "text", "text": paragraph.text}]
    
    result = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        
        # Build text node with marks for formatting
        marks = []
        if run.bold:
            marks.append({"type": "bold"})
        if run.italic:
            marks.append({"type": "italic"})
        if run.underline:
            marks.append({"type": "underline"})
        
        node = {"type": "text", "text": text}
        if marks:
            node["marks"] = marks
        result.append(node)
    
    return result if result else [{"type": "text", "text": paragraph.text}]

def get_list_info(paragraph):
    """Detect if paragraph is a list item and get its properties"""
    try:
        # Check paragraph XML for list properties
        pPr = paragraph._element.pPr
        if pPr is not None:
            numPr = pPr.numPr
            if numPr is not None:
                ilvl = numPr.ilvl
                level = int(ilvl.val) if ilvl is not None else 0
                
                numId = numPr.numId
                num_id = int(numId.val) if numId is not None else 0
                
                # Determine list type (bullet vs ordered)
                # Numbers 1-9 are typically bullets, 10+ are numbered
                is_ordered = num_id >= 10 or "number" in (paragraph.style.name.lower() if paragraph.style else "")
                
                return {"is_list": True, "level": level, "ordered": is_ordered, "num_id": num_id}
    except Exception:
        pass
    
    # Fallback: check style name
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    if "list" in style_name:
        is_ordered = "number" in style_name or "ordered" in style_name
        return {"is_list": True, "level": 0, "ordered": is_ordered, "num_id": 0}
    
    # Fallback: check text prefix
    text = paragraph.text.strip()
    if text.startswith(('•', '●', '○', '▪', '-', '*')):
        return {"is_list": True, "level": 0, "ordered": False, "num_id": -1}
    if len(text) > 2 and text[0].isdigit() and text[1] in '.):':
        return {"is_list": True, "level": 0, "ordered": True, "num_id": -2}
    
    return {"is_list": False}

def extract_table(table):
    """Extract table structure from DOCX table"""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            # Get cell text (join all paragraphs)
            cell_text = "\n".join(p.text for p in cell.paragraphs)
            cells.append({"type": "tableCell", "children": [{"type": "text", "text": cell_text}]})
        rows.append({"type": "tableRow", "children": cells})
    
    return {"type": "table", "children": rows}

def extract_text_from_docx(minio, bucket_name, filename):
    """Extract structured content from a DOCX file with full formatting support"""
    print(f"[worker] Extracting text from DOCX: {bucket_name}/{filename}")
    
    try:
        # Download DOCX from MinIO
        response = minio.get_object(bucket_name, filename)
        docx_bytes = response.read()
        response.close()
        response.release_conn()
        
        # Open DOCX with python-docx
        doc = DocxDocument(BytesIO(docx_bytes))
        
        content_blocks = []
        current_list = None
        current_list_type = None
        
        # Process document body - paragraphs and tables in order
        for element in doc.element.body:
            # Check if it's a table
            if element.tag.endswith('tbl'):
                # Flush any current list
                if current_list:
                    content_blocks.append(current_list)
                    current_list = None
                
                # Find the table object
                for table in doc.tables:
                    if table._element == element:
                        content_blocks.append({
                            "type": "table",
                            "rows": [[cell.text for cell in row.cells] for row in table.rows]
                        })
                        break
            
            # Check if it's a paragraph
            elif element.tag.endswith('p'):
                # Find matching paragraph object
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                
                if para is None:
                    continue
                
                text = para.text.strip()
                if not text:
                    # Empty paragraph - flush list
                    if current_list:
                        content_blocks.append(current_list)
                        current_list = None
                    continue
                
                # Check for heading
                style_name = para.style.name.lower() if para.style else ""
                
                if "heading" in style_name or "title" in style_name or "kop" in style_name:
                    # Flush list
                    if current_list:
                        content_blocks.append(current_list)
                        current_list = None
                    
                    # Extract heading level
                    level = 1
                    for char in style_name:
                        if char.isdigit():
                            level = int(char)
                            break
                    
                    content_blocks.append({
                        "type": "heading",
                        "level": level,
                        "text": text,
                        "children": extract_runs_with_formatting(para)
                    })
                    continue
                
                # Check for list
                list_info = get_list_info(para)
                if list_info["is_list"]:
                    list_type = "orderedList" if list_info["ordered"] else "bulletList"
                    
                    # Start new list or continue existing
                    if current_list is None or current_list_type != list_type:
                        if current_list:
                            content_blocks.append(current_list)
                        current_list = {"type": list_type, "items": []}
                        current_list_type = list_type
                    
                    current_list["items"].append({
                        "type": "listItem",
                        "text": text,
                        "children": extract_runs_with_formatting(para)
                    })
                    continue
                
                # Regular paragraph - flush any list first
                if current_list:
                    content_blocks.append(current_list)
                    current_list = None
                
                # Check if bold text might be a heading
                is_bold_heading = False
                if para.runs:
                    first_run = para.runs[0]
                    if first_run.bold and len(text.split()) < 15:
                        # Check font size too
                        if first_run.font.size and first_run.font.size >= Pt(14):
                            is_bold_heading = True
                        elif first_run.bold and len(text.split()) < 8:
                            is_bold_heading = True
                
                if is_bold_heading:
                    content_blocks.append({
                        "type": "heading",
                        "level": 3,
                        "text": text,
                        "children": extract_runs_with_formatting(para)
                    })
                else:
                    content_blocks.append({
                        "type": "paragraph",
                        "text": text,
                        "children": extract_runs_with_formatting(para)
                    })
        
        # Flush any remaining list
        if current_list:
            content_blocks.append(current_list)
        
        print(f"[worker] DOCX: {len(content_blocks)} content blocks (including tables/lists)")
        
        # Count different types
        types_count = {}
        for block in content_blocks:
            t = block.get("type", "unknown")
            types_count[t] = types_count.get(t, 0) + 1
        print(f"[worker] Block types: {types_count}")
        
        # Return as single "page"
        return [{"page_number": 1, "content": content_blocks}]
        
    except Exception as e:
        print(f"[worker] Error extracting DOCX text: {e}")
        import traceback
        traceback.print_exc()
        return None

def detect_file_type(minio, bucket_name, filename):
    """Detect file type from content (magic bytes)"""
    try:
        response = minio.get_object(bucket_name, filename, length=8)
        header = response.read()
        response.close()
        response.release_conn()
        
        # PDF starts with %PDF
        if header.startswith(b'%PDF'):
            return 'pdf'
        
        # DOCX (ZIP with specific content) starts with PK
        if header.startswith(b'PK\x03\x04'):
            return 'docx'
        
        return 'unknown'
    except Exception as e:
        print(f"[worker] Error detecting file type: {e}")
        return 'unknown'

def fix_encoding(text):
    """Fix common encoding issues from PDF extraction"""
    # Common Windows-1252 to UTF-8 misinterpretations
    replacements = {
        'â€"': '—',  # em-dash
        'â€"': '–',  # en-dash
        'â€™': "'",  # right single quote
        'â€œ': '"',  # left double quote
        'â€': '"',   # right double quote
        'â€¦': '…',  # ellipsis
        'Ã©': 'é',
        'Ã¨': 'è',
        'Ã«': 'ë',
        'Ã¯': 'ï',
        'Ã¶': 'ö',
        'Ã¼': 'ü',
        'Ã ': 'à',
        'ΓÇô': '–',  # en-dash variant
        'ΓÇö': '—',  # em-dash variant
        'ΓÇï': '',   # zero-width space
        'ΓÇ£': '"',  # quote
        'ΓÇª': '…',  # ellipsis
        '\u200b': '', # zero-width space
        '\ufeff': '', # BOM
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    return text

def make_headers(trace_id):
    return {
        "x-kimi-worker-instance-name": INSTANCE_NAME,
        "x-kimi-worker-name": WORKER_NAME,
        "x-trace-id": trace_id,
        "timestamp": datetime.utcnow().isoformat()
    }

def extract_page_content(attributes):
    """Extract page content from job attributes"""
    pages = []
    
    # Check if we have content attribute
    content_attr = attributes.get("content", {})
    if "values" in content_attr:
        for value in content_attr["values"]:
            if "stringResult" in value:
                try:
                    page_content = json.loads(value["stringResult"])
                    pages.append(page_content)
                except json.JSONDecodeError:
                    pages.append({"text": value["stringResult"]})
    
    return pages

def convert_text_children_to_spec(text_children):
    """Convert text children with marks to NLdoc spec format"""
    result = []
    for child in text_children:
        if child.get("type") == "text":
            text_node = {
                "id": str(uuid.uuid4()),
                "type": "https://spec.nldoc.nl/Resource/Text",
                "text": child.get("text", "")
            }
            # Add marks if present (bold, italic, underline)
            marks = child.get("marks", [])
            if marks:
                text_node["marks"] = marks
            result.append(text_node)
    return result if result else [{"id": str(uuid.uuid4()), "type": "https://spec.nldoc.nl/Resource/Text", "text": ""}]

def generate_spec_from_content(doc_id, page_count, page_contents):
    """Generate NLdoc spec JSON from structured page contents
    
    NLdoc spec format uses https://spec.nldoc.nl/Resource/* types:
    - Document: root element
    - Heading: heading with level (10=h1, 20=h2, 30=h3, etc)
    - Paragraph: paragraph container
    - Text: text content with optional marks
    - Table: table container
    - TableRow: table row
    - TableCell: table cell
    - BulletList/OrderedList: list containers
    - ListItem: list item
    """
    
    children = []
    
    if page_contents:
        for page in page_contents:
            content_blocks = page.get("content", [])
            
            for block in content_blocks:
                block_type = block.get("type", "paragraph")
                text = block.get("text", "").strip()
                text_children = block.get("children", [])
                
                # Handle headings
                if block_type == "heading":
                    level = block.get("level", 2)
                    # NLdoc uses level * 10 (10=h1, 20=h2, 30=h3, etc)
                    heading_level = level * 10 if level <= 6 else 60
                    
                    spec_children = convert_text_children_to_spec(text_children) if text_children else [{
                        "id": str(uuid.uuid4()),
                        "type": "https://spec.nldoc.nl/Resource/Text",
                        "text": text
                    }]
                    
                    children.append({
                        "id": str(uuid.uuid4()),
                        "type": "https://spec.nldoc.nl/Resource/Heading",
                        "level": heading_level,
                        "children": spec_children
                    })
                
                # Handle tables
                elif block_type == "table":
                    rows = block.get("rows", [])
                    table_rows = []
                    
                    for row_idx, row in enumerate(rows):
                        table_cells = []
                        for cell_text in row:
                            table_cells.append({
                                "id": str(uuid.uuid4()),
                                "type": "https://spec.nldoc.nl/Resource/TableCell",
                                "children": [{
                                    "id": str(uuid.uuid4()),
                                    "type": "https://spec.nldoc.nl/Resource/Paragraph",
                                    "children": [{
                                        "id": str(uuid.uuid4()),
                                        "type": "https://spec.nldoc.nl/Resource/Text",
                                        "text": cell_text
                                    }]
                                }]
                            })
                        
                        # First row might be header
                        row_type = "https://spec.nldoc.nl/Resource/TableHeaderRow" if row_idx == 0 else "https://spec.nldoc.nl/Resource/TableRow"
                        table_rows.append({
                            "id": str(uuid.uuid4()),
                            "type": row_type,
                            "children": table_cells
                        })
                    
                    if table_rows:
                        children.append({
                            "id": str(uuid.uuid4()),
                            "type": "https://spec.nldoc.nl/Resource/Table",
                            "children": table_rows
                        })
                
                # Handle bullet lists
                elif block_type == "bulletList":
                    list_items = []
                    for item in block.get("items", []):
                        item_text = item.get("text", "")
                        item_children = item.get("children", [])
                        
                        spec_children = convert_text_children_to_spec(item_children) if item_children else [{
                            "id": str(uuid.uuid4()),
                            "type": "https://spec.nldoc.nl/Resource/Text",
                            "text": item_text
                        }]
                        
                        list_items.append({
                            "id": str(uuid.uuid4()),
                            "type": "https://spec.nldoc.nl/Resource/ListItem",
                            "children": [{
                                "id": str(uuid.uuid4()),
                                "type": "https://spec.nldoc.nl/Resource/Paragraph",
                                "children": spec_children
                            }]
                        })
                    
                    if list_items:
                        children.append({
                            "id": str(uuid.uuid4()),
                            "type": "https://spec.nldoc.nl/Resource/BulletList",
                            "children": list_items
                        })
                
                # Handle ordered lists
                elif block_type == "orderedList":
                    list_items = []
                    for idx, item in enumerate(block.get("items", [])):
                        item_text = item.get("text", "")
                        item_children = item.get("children", [])
                        
                        spec_children = convert_text_children_to_spec(item_children) if item_children else [{
                            "id": str(uuid.uuid4()),
                            "type": "https://spec.nldoc.nl/Resource/Text",
                            "text": item_text
                        }]
                        
                        list_items.append({
                            "id": str(uuid.uuid4()),
                            "type": "https://spec.nldoc.nl/Resource/ListItem",
                            "order": idx + 1,
                            "children": [{
                                "id": str(uuid.uuid4()),
                                "type": "https://spec.nldoc.nl/Resource/Paragraph",
                                "children": spec_children
                            }]
                        })
                    
                    if list_items:
                        children.append({
                            "id": str(uuid.uuid4()),
                            "type": "https://spec.nldoc.nl/Resource/OrderedList",
                            "children": list_items
                        })
                
                # Handle paragraphs
                elif block_type == "paragraph":
                    if not text:
                        continue
                    
                    spec_children = convert_text_children_to_spec(text_children) if text_children else [{
                        "id": str(uuid.uuid4()),
                        "type": "https://spec.nldoc.nl/Resource/Text",
                        "text": text
                    }]
                    
                    children.append({
                        "id": str(uuid.uuid4()),
                        "type": "https://spec.nldoc.nl/Resource/Paragraph",
                        "children": spec_children
                    })
    
    if not children:
        # Fallback: no content available
        children.append({
            "id": str(uuid.uuid4()),
            "type": "https://spec.nldoc.nl/Resource/Paragraph",
            "children": [{
                "id": str(uuid.uuid4()),
                "type": "https://spec.nldoc.nl/Resource/Text",
                "text": f"Dit document bevat {page_count} pagina's maar de tekst kon niet worden geëxtraheerd."
            }]
        })
    
    return {
        "id": str(uuid.uuid4()),
        "type": "https://spec.nldoc.nl/Resource/Document",
        "children": children
    }

def spec_to_html(spec):
    """Convert NLdoc spec JSON to accessible HTML"""
    
    def render_marks(text, marks):
        """Apply marks (bold, italic, underline) to text"""
        if not marks:
            return html_escape(text)
        
        result = html_escape(text)
        for mark in marks:
            mark_type = mark.get("type", "")
            if mark_type == "bold" or mark_type == "strong":
                result = f"<strong>{result}</strong>"
            elif mark_type == "italic" or mark_type == "em":
                result = f"<em>{result}</em>"
            elif mark_type == "underline":
                result = f"<u>{result}</u>"
        return result
    
    def html_escape(text):
        """Escape HTML special characters"""
        if not text:
            return ""
        return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#x27;"))
    
    def render_children(children):
        """Render an array of children nodes"""
        result = []
        for child in children:
            result.append(render_node(child))
        return "".join(result)
    
    def render_node(node):
        """Render a single node to HTML"""
        node_type = node.get("type", "")
        children = node.get("children", [])
        
        # Text node
        if node_type.endswith("/Text"):
            text = node.get("text", "")
            marks = node.get("marks", [])
            return render_marks(text, marks)
        
        # Heading
        if node_type.endswith("/Heading"):
            level = node.get("level", 20)
            h_level = min(6, max(1, level // 10))  # Convert 10->1, 20->2, etc
            content = render_children(children)
            return f"<h{h_level}>{content}</h{h_level}>\n"
        
        # Paragraph
        if node_type.endswith("/Paragraph"):
            content = render_children(children)
            return f"<p>{content}</p>\n"
        
        # BulletList
        if node_type.endswith("/BulletList"):
            content = render_children(children)
            return f"<ul>\n{content}</ul>\n"
        
        # OrderedList
        if node_type.endswith("/OrderedList"):
            content = render_children(children)
            return f"<ol>\n{content}</ol>\n"
        
        # ListItem
        if node_type.endswith("/ListItem"):
            content = render_children(children)
            return f"<li>{content}</li>\n"
        
        # Table
        if node_type.endswith("/Table"):
            content = render_children(children)
            return f"<table>\n{content}</table>\n"
        
        # TableHeaderRow
        if node_type.endswith("/TableHeaderRow"):
            cells = []
            for child in children:
                cell_content = render_children(child.get("children", []))
                cells.append(f"<th>{cell_content}</th>")
            return f"<tr>{''.join(cells)}</tr>\n"
        
        # TableRow
        if node_type.endswith("/TableRow"):
            cells = []
            for child in children:
                cell_content = render_children(child.get("children", []))
                cells.append(f"<td>{cell_content}</td>")
            return f"<tr>{''.join(cells)}</tr>\n"
        
        # TableCell (standalone, shouldn't normally be rendered directly)
        if node_type.endswith("/TableCell"):
            content = render_children(children)
            return content
        
        # Document
        if node_type.endswith("/Document"):
            return render_children(children)
        
        # Unknown - try to render children
        if children:
            return render_children(children)
        
        return ""
    
    # Generate HTML
    body = render_node(spec)
    
    html = f'''<!DOCTYPE html>
<html lang="nl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Geconverteerd Document</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 2rem; color: #333; }}
        h1, h2, h3, h4, h5, h6 {{ margin-top: 1.5em; margin-bottom: 0.5em; color: #1a1a1a; }}
        h1 {{ font-size: 2rem; border-bottom: 2px solid #eee; padding-bottom: 0.3em; }}
        h2 {{ font-size: 1.5rem; border-bottom: 1px solid #eee; padding-bottom: 0.2em; }}
        h3 {{ font-size: 1.25rem; }}
        p {{ margin: 1em 0; }}
        ul, ol {{ margin: 1em 0; padding-left: 2em; }}
        li {{ margin: 0.3em 0; }}
        table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
        th, td {{ border: 1px solid #ddd; padding: 0.75em; text-align: left; }}
        th {{ background-color: #f5f5f5; font-weight: bold; }}
        tr:nth-child(even) {{ background-color: #fafafa; }}
        strong {{ font-weight: bold; }}
        em {{ font-style: italic; }}
        u {{ text-decoration: underline; }}
    </style>
</head>
<body>
{body}
</body>
</html>'''
    
    return html

def spec_to_tiptap(spec):
    """Convert NLdoc spec JSON (Resource/*) to TipTap JSON (doc/content)."""

    def marks_to_tiptap(marks):
        # NLdoc marks are stored as [{"type":"bold"|"italic"|"underline"|...}]
        if not marks:
            return None
        out = []
        for m in marks:
            t = (m or {}).get("type")
            if t in ("bold", "strong"):
                out.append({"type": "bold"})
            elif t in ("italic", "em"):
                out.append({"type": "italic"})
            elif t == "underline":
                out.append({"type": "underline"})
        return out or None

    def as_list(v):
        return v if isinstance(v, list) else []

    def render_inline(children):
        # Returns list of tiptap inline nodes
        out = []
        for c in as_list(children):
            t = c.get("type", "")
            if t.endswith("/Text"):
                node = {"type": "text", "text": c.get("text", "")}
                marks = marks_to_tiptap(c.get("marks"))
                if marks:
                    node["marks"] = marks
                out.append(node)
        return out

    def render_block(node):
        t = node.get("type", "")
        children = node.get("children", [])

        if t.endswith("/Heading"):
            level = int(node.get("level", 20) // 10) if isinstance(node.get("level"), int) else 2
            level = max(1, min(6, level))
            return {
                "type": "heading",
                "attrs": {"level": level},
                "content": render_inline(children) or [{"type": "text", "text": ""}],
            }

        if t.endswith("/Paragraph"):
            return {
                "type": "paragraph",
                "content": render_inline(children) or [{"type": "text", "text": ""}],
            }

        if t.endswith("/BulletList"):
            return {
                "type": "bulletList",
                "content": [render_block(c) for c in as_list(children) if render_block(c)],
            }

        if t.endswith("/OrderedList"):
            return {
                "type": "orderedList",
                "content": [render_block(c) for c in as_list(children) if render_block(c)],
            }

        if t.endswith("/ListItem"):
            # children are typically Paragraph nodes
            return {
                "type": "listItem",
                "content": [render_block(c) for c in as_list(children) if render_block(c)],
            }

        if t.endswith("/Table"):
            rows = []
            for r in as_list(children):
                rt = r.get("type", "")
                rchildren = as_list(r.get("children", []))
                if rt.endswith("/TableHeaderRow"):
                    cells = []
                    for cell in rchildren:
                        # NLdoc: TableCell -> [Paragraph...]
                        cell_paras = [render_block(p) for p in as_list(cell.get("children", [])) if render_block(p)]
                        cells.append({"type": "tableHeader", "content": cell_paras or [{"type": "paragraph", "content": [{"type": "text", "text": ""}]}]})
                    rows.append({"type": "tableRow", "content": cells})
                elif rt.endswith("/TableRow"):
                    cells = []
                    for cell in rchildren:
                        cell_paras = [render_block(p) for p in as_list(cell.get("children", [])) if render_block(p)]
                        cells.append({"type": "tableCell", "content": cell_paras or [{"type": "paragraph", "content": [{"type": "text", "text": ""}]}]})
                    rows.append({"type": "tableRow", "content": cells})
            return {"type": "table", "content": rows}

        # Unknown/unsupported block
        return None

    doc_children = []
    if isinstance(spec, dict) and str(spec.get("type", "")).endswith("/Document"):
        for c in as_list(spec.get("children", [])):
            b = render_block(c)
            if b:
                doc_children.append(b)

    return {"type": "doc", "content": doc_children}

def send_worker_result(ch, doc_id, job_id, spec, trace_id, job_metadata):
    """Send worker result back to station-folio-spec and generate HTML"""
    
    now = datetime.utcnow().isoformat() + "Z"
    minio = get_minio()
    
    # Ensure buckets exist
    if not minio.bucket_exists("files"):
        minio.make_bucket("files")
    if not minio.bucket_exists("output"):
        minio.make_bucket("output")
    
    # Upload spec to MinIO
    spec_json = json.dumps(spec, ensure_ascii=False)
    spec_bytes = spec_json.encode("utf-8")
    minio.put_object("files", f"{doc_id}.spec.json", BytesIO(spec_bytes), len(spec_bytes), content_type="application/json")
    print(f"[worker] Uploaded spec to files/{doc_id}.spec.json")
    
    # Generate and upload HTML directly (bypassing html-writer)
    html_content = spec_to_html(spec)
    html_bytes = html_content.encode("utf-8")
    minio.put_object(
        "output",
        f"{doc_id}.html",
        BytesIO(html_bytes),
        len(html_bytes),
        content_type="text/html; charset=utf-8",
    )
    print(f"[worker] Uploaded HTML to output/{doc_id}.html ({len(html_bytes)} bytes)")

    # If the requested target is TipTap JSON, also generate & upload that and point DONE to it.
    target = job_metadata.get("targetFileType", "text/html")
    tiptap_location = None
    if target == "application/vnd.nldoc.tiptap+json":
        tiptap_doc = spec_to_tiptap(spec)
        tiptap_bytes = json.dumps(tiptap_doc, ensure_ascii=False).encode("utf-8")
        minio.put_object(
            "output",
            f"{doc_id}.json",
            BytesIO(tiptap_bytes),
            len(tiptap_bytes),
            content_type="application/vnd.nldoc.tiptap+json; charset=utf-8",
        )
        tiptap_location = f"{doc_id}.json"
        print(f"[worker] Uploaded TipTap JSON to output/{doc_id}.json ({len(tiptap_bytes)} bytes)")
    
    result = {
        "resultType": "fileWorkerResult",
        "traceId": trace_id,
        "recordId": f"folio|||{doc_id}",
        "jobId": job_id,
        "timestamp": now,
        "success": True,
        "bucketName": "files",
        "filename": f"{doc_id}.spec.json"
    }
    
    # Publish result to station-folio-spec
    headers = make_headers(trace_id)
    ch.basic_publish(
        exchange="nldoc.topics",
        routing_key=f"worker.folio-spec.results.{job_id}",
        body=json.dumps(result),
        properties=pika.BasicProperties(
            content_type="application/json",
            headers=headers
        )
    )
    print(f"[worker] Published result to worker.folio-spec.results.{job_id}")
    
    # Publish DONE event for the API/editor stream.
    # The editor expects the event-types schema:
    #   type = https://event.spec.nldoc.nl/done
    #   context.location = "<filename>" (without /file prefix)
    #   context.contentType = requested target content-type
    done_message = {
        "type": "https://event.spec.nldoc.nl/done",
        "timestamp": now,
        "traceId": trace_id,
        "context": {
            "contentType": target,
            "location": tiptap_location or f"{doc_id}.html",
        },
    }

    # Give clients a small window to subscribe to SSE after upload, otherwise the DONE event
    # can be published before the client is listening (the API stream does not always replay history).
    try:
        delay_s = float(os.environ.get("DONE_DELAY_SECONDS", "2"))
        if delay_s > 0:
            time.sleep(delay_s)
    except Exception:
        pass
    
    # IMPORTANT: do NOT publish this to documents.{docId}. That topic is consumed by station-document-source
    # which expects a DocumentInfoReport payload and will emit https://event.spec.nldoc.nl/error if parsing fails.
    ch.basic_publish(
        exchange="nldoc.topics",
        routing_key=f"events.{doc_id}",
        body=json.dumps(done_message),
        properties=pika.BasicProperties(
            content_type="application/json",
            headers=headers,
            delivery_mode=2
        )
    )
    print(f"[worker] Published done event to events.{doc_id}")

def handle_job(ch, method, properties, body):
    """Process a worker job from folio-spec station"""
    try:
        job = json.loads(body)
        routing_key = method.routing_key
        
        print(f"[worker] Received job on {routing_key}")
        print(f"[worker] Job keys: {list(job.keys())}")
        
        # Extract job info
        record_id = job.get("recordId", "")
        doc_id = record_id.split("|||")[-1] if "|||" in record_id else job.get("filename", "unknown")
        job_id = job.get("jobId", str(uuid.uuid4()))
        bucket_name = job.get("bucketName", "files")
        target_file_type = job.get("targetFileType", "text/html")
        
        # Get trace ID from headers or job
        trace_id = doc_id
        if properties.headers:
            trace_id = properties.headers.get("x-trace-id", doc_id)
        
        print(f"[worker] Processing doc: {doc_id}, target: {target_file_type}")
        
        # Extract attributes
        attributes = job.get("attributes", {})
        
        # Get page count
        page_count = 10  # default
        page_count_attr = attributes.get("pageCount", {})
        if "values" in page_count_attr and len(page_count_attr["values"]) > 0:
            page_count = int(page_count_attr["values"][0].get("stringResult", 10))
        
        print(f"[worker] Page count: {page_count}")
        
        # Extract text from document (PDF or DOCX)
        minio = get_minio()
        original_filename = job.get("filename", doc_id)
        
        # Detect file type
        file_type = detect_file_type(minio, bucket_name, original_filename)
        print(f"[worker] Detected file type: {file_type}")
        
        page_contents = None
        if file_type == 'pdf':
            page_contents = extract_text_from_pdf(minio, bucket_name, original_filename)
        elif file_type == 'docx':
            page_contents = extract_text_from_docx(minio, bucket_name, original_filename)
        else:
            print(f"[worker] Unknown file type, trying PDF first then DOCX")
            page_contents = extract_text_from_pdf(minio, bucket_name, original_filename)
            if not page_contents:
                page_contents = extract_text_from_docx(minio, bucket_name, original_filename)
        
        if page_contents:
            print(f"[worker] Extracted text from {len(page_contents)} pages/sections")
        else:
            print(f"[worker] Could not extract text, using fallback")
            page_contents = []
        
        # Generate spec
        spec = generate_spec_from_content(doc_id, page_count, page_contents)
        
        # Collect job metadata for downstream stations
        job_metadata = {
            "bucketName": bucket_name,
            "filename": job.get("filename", doc_id),
            "targetFileType": target_file_type,
            "kimiRegistrationDate": job.get("kimiRegistrationDate", datetime.utcnow().isoformat() + "Z"),
            "creationDate": job.get("creationDate", datetime.utcnow().isoformat() + "Z"),
            "processCount": job.get("processCount", 0),
        }
        
        # Send result
        send_worker_result(ch, doc_id, job_id, spec, trace_id, job_metadata)
        
        # ACK the message
        ch.basic_ack(delivery_tag=method.delivery_tag)
        print(f"[worker] Completed job for {doc_id}")
        
    except Exception as e:
        print(f"[worker] Error processing job: {e}")
        import traceback
        traceback.print_exc()
        ch.basic_nack(delivery_tag=method.delivery_tag, requeue=False)

def send_heartbeat(ch):
    """Send worker heartbeat"""
    # Must match the station WorkerAvailability schema (zod):
    # - workerType: string
    # - workerInstance: string
    # - state: "active" | "interrupted" | "stopped"
    # - reason: string
    now = datetime.utcnow().isoformat() + "Z"
    heartbeat = {
        "workerType": WORKER_NAME,
        "workerName": "folio-spec-worker",
        "workerInstance": INSTANCE_NAME,
        "state": "active",
        "reason": "health report",
        "timestamp": now,
    }
    ch.basic_publish(
        exchange="nldoc.topics",
        routing_key=f"worker.folio-spec.health.heartbeats.{INSTANCE_NAME}",
        body=json.dumps(heartbeat),
        properties=pika.BasicProperties(content_type="application/json")
    )

def main():
    import threading

    print(f"[worker] Starting {INSTANCE_NAME}...")
    print(f"[worker] AMQP_URL: {AMQP_URL[:50] if AMQP_URL else 'None'}...")

    # Robust consume loop: pika can drop connections (e.g. laptop sleep / network blip).
    # We reconnect instead of crashing the container (CrashLoopBackOff).
    while True:
        stop_hb = threading.Event()
        conn = None

        try:
            params = pika.URLParameters(AMQP_URL)
            params.heartbeat = 60
            params.blocked_connection_timeout = 300

            conn = pika.BlockingConnection(params)
            ch = conn.channel()

            # Declare our queue
            ch.queue_declare(queue="worker-folio-spec", durable=True)

            # Bind to worker.folio-spec.jobs (for PDF flow)
            ch.queue_bind(
                exchange="nldoc.topics",
                queue="worker-folio-spec",
                routing_key="worker.folio-spec.jobs",
            )

            # Also bind to worker.docx-spec.jobs (for DOCX flow)
            ch.queue_bind(
                exchange="nldoc.topics",
                queue="worker-folio-spec",
                routing_key="worker.docx-spec.jobs",
            )

            print("[worker] Bound to worker.folio-spec.jobs AND worker.docx-spec.jobs")

            # Send initial heartbeat
            send_heartbeat(ch)

            # Heartbeat thread
            def heartbeat_loop():
                while not stop_hb.is_set():
                    time.sleep(30)
                    try:
                        send_heartbeat(ch)
                    except Exception:
                        # Connection/channel might be gone; reconnect loop will handle.
                        pass

            threading.Thread(target=heartbeat_loop, daemon=True).start()

            # Set QoS and start consuming
            ch.basic_qos(prefetch_count=1)
            ch.basic_consume(queue="worker-folio-spec", on_message_callback=handle_job)
            print("[worker] Waiting for jobs...")
            ch.start_consuming()

        except Exception as e:
            print(f"[worker] Connection/consume error: {e}")
            import traceback
            traceback.print_exc()
            time.sleep(5)

        finally:
            stop_hb.set()
            try:
                if conn and conn.is_open:
                    conn.close()
            except Exception:
                pass

if __name__ == "__main__":
    main()

